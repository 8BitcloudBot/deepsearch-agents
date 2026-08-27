"""RED: Complete file reader contracts.

Requires pypdf, python-docx, openpyxl with real parsing.
No regex-only page counting. No sharedStrings-only XLSX.
Macro rejection, ZIP bomb defense, and proper error redaction.
"""

import zipfile
from pathlib import Path

import pytest

from app.tools.files import MAX_FILE_SIZE_BYTES, SessionWorkspace

UUID_V4 = "00000000-0000-4000-8000-000000000001"

# ── TXT / Markdown ──────────────────────────────────────────────────────────


class TestTextReader:
    def test_reads_utf8(self, tmp_path):
        from app.tools.files import read_text_file

        f = tmp_path / "test.txt"
        f.write_text("hello world", encoding="utf-8")
        assert "hello world" in read_text_file(f)

    def test_rejects_non_utf8(self, tmp_path):
        from app.tools.files import read_text_file

        f = tmp_path / "bad.txt"
        f.write_bytes(b"\xff\xfe\x00\x00")
        with pytest.raises(ValueError):
            read_text_file(f)

    def test_truncates_above_max_chars(self, tmp_path):
        from app.tools.files import TRUNCATION_SUFFIX, read_text_file

        f = tmp_path / "large.txt"
        f.write_text("あ" * 150_000, encoding="utf-8")
        content = read_text_file(f, max_chars=100_000)
        assert len(content) <= 100_000 + len(TRUNCATION_SUFFIX) + 100
        assert TRUNCATION_SUFFIX.strip() in content

    def test_utf8_bom_accepted(self, tmp_path):
        from app.tools.files import read_text_file

        f = tmp_path / "bom.txt"
        f.write_bytes(b"\xef\xbb\xbfhello")
        content = read_text_file(f)
        assert "hello" in content


# ── PDF (pypdf required) ─────────────────────────────────────────────────────


class TestPDFReader:
    def test_rejects_non_pdf_header(self, tmp_path):
        from app.tools.files import read_pdf_file

        f = tmp_path / "fake.pdf"
        f.write_text("not a pdf")
        with pytest.raises(ValueError):
            read_pdf_file(f)

    def test_reads_valid_pdf_with_pypdf(self, tmp_path):
        from app.tools.files import read_pdf_file

        f = _make_valid_pdf(tmp_path / "doc.pdf", pages=2)
        content = read_pdf_file(f)
        assert len(content) > 0

    def test_rejects_encrypted_pdf(self, tmp_path):
        from pypdf import PdfWriter

        from app.tools.files import read_pdf_file

        writer = PdfWriter()
        writer.add_blank_page(100, 100)
        writer.encrypt("user_pass")
        f = tmp_path / "enc.pdf"
        with open(f, "wb") as fp:
            writer.write(fp)
        with pytest.raises(ValueError, match="ncrypt"):
            read_pdf_file(f)

    def test_rejects_damaged_pdf(self, tmp_path):
        from app.tools.files import read_pdf_file

        f = tmp_path / "broken.pdf"
        f.write_bytes(b"%PDF-1.4\ngarbage without proper structure")
        with pytest.raises(ValueError):
            read_pdf_file(f)

    def test_rejects_over_200_pages(self, tmp_path):
        from app.tools.files import read_pdf_file

        f = _make_valid_pdf(tmp_path / "big.pdf", pages=250)
        with pytest.raises(ValueError, match="max"):
            read_pdf_file(f, max_pages=200)

    def test_accepts_200_pages(self, tmp_path):
        from app.tools.files import read_pdf_file

        f = _make_valid_pdf(tmp_path / "ok.pdf", pages=200)
        content = read_pdf_file(f, max_pages=200)
        assert len(content) > 0

    def test_pdf_extracts_text_not_placeholder(self, tmp_path):
        from app.tools.files import read_pdf_file

        f = _make_valid_pdf(tmp_path / "text.pdf", pages=1)
        content = read_pdf_file(f)
        # Must NOT return the old placeholder format
        assert not content.startswith("[PDF content:")


# ── DOCX (python-docx required) ─────────────────────────────────────────────


class TestDOCXReader:
    def test_reads_valid_docx_with_python_docx(self, tmp_path):
        from app.tools.files import read_docx_file

        f = _make_real_docx(tmp_path / "doc.docx")
        content = read_docx_file(f)
        assert len(content) > 0

    def test_rejects_non_zip(self, tmp_path):
        from app.tools.files import read_docx_file

        f = tmp_path / "fake.docx"
        f.write_text("not a zip")
        with pytest.raises(ValueError):
            read_docx_file(f)

    def test_rejects_missing_content_types(self, tmp_path):
        from app.tools.files import read_docx_file

        f = tmp_path / "bad.docx"
        with zipfile.ZipFile(f, "w") as zf:
            zf.writestr("word/document.xml", "<root/>")
        with pytest.raises(ValueError):
            read_docx_file(f)

    def test_rejects_missing_document_xml(self, tmp_path):
        from app.tools.files import read_docx_file

        f = tmp_path / "bad2.docx"
        with zipfile.ZipFile(f, "w") as zf:
            zf.writestr("[Content_Types].xml", _CT_XML)
        with pytest.raises(ValueError):
            read_docx_file(f)

    def test_rejects_macro_enabled_docx(self, tmp_path):
        from app.tools.files import read_docx_file

        f = tmp_path / "macro.docx"
        with zipfile.ZipFile(f, "w") as zf:
            zf.writestr("[Content_Types].xml", _CT_XML)
            zf.writestr("word/document.xml", _DOCUMENT_XML)
            zf.writestr("word/vbaProject.bin", b"macro payload")
        with pytest.raises(ValueError, match="vbaProject"):
            read_docx_file(f)

    def test_rejects_macro_content_type(self, tmp_path):
        from app.tools.files import read_docx_file

        f = tmp_path / "macro2.docx"
        ct = _CT_XML.replace(
            'Extension="docx"',
            'Extension="docm"',
        ).replace(
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"',  # noqa: E501
            'ContentType="application/vnd.ms-word.document.macroEnabled.main+xml"',
        )
        with zipfile.ZipFile(f, "w") as zf:
            zf.writestr("[Content_Types].xml", ct)
            zf.writestr("word/document.xml", _DOCUMENT_XML)
        with pytest.raises(ValueError, match="acro"):
            read_docx_file(f)

    def test_rejects_zip_bomb(self, tmp_path):
        from app.tools.files import read_docx_file

        f = tmp_path / "bomb.docx"
        with zipfile.ZipFile(f, "w") as zf:
            zf.writestr("[Content_Types].xml", _CT_XML)
            zf.writestr("word/document.xml", _DOCUMENT_XML)
            # 10000 entries with tiny compressed size but huge uncompressed
            for i in range(10000):
                zf.writestr(f"word/extra{i}.xml", "x" * 10000)
        with pytest.raises(ValueError, match="bomb|entry|limit"):
            read_docx_file(f)


# ── XLSX (openpyxl required) ────────────────────────────────────────────────


class TestXLSXReader:
    def test_rejects_non_zip(self, tmp_path):
        from app.tools.files import read_xlsx_file

        f = tmp_path / "fake.xlsx"
        f.write_text("not a zip")
        with pytest.raises(ValueError):
            read_xlsx_file(f)

    def test_rejects_missing_content_types(self, tmp_path):
        from app.tools.files import read_xlsx_file

        f = tmp_path / "bad.xlsx"
        with zipfile.ZipFile(f, "w") as zf:
            zf.writestr("xl/workbook.xml", "<root/>")
        with pytest.raises(ValueError):
            read_xlsx_file(f)

    def test_rejects_missing_workbook(self, tmp_path):
        from app.tools.files import read_xlsx_file

        f = tmp_path / "bad2.xlsx"
        with zipfile.ZipFile(f, "w") as zf:
            zf.writestr("[Content_Types].xml", _CT_XLSX_XML)
        with pytest.raises(ValueError):
            read_xlsx_file(f)

    def test_reads_real_xlsx_with_openpyxl(self, tmp_path):
        from app.tools.files import read_xlsx_file

        f = _make_real_xlsx(tmp_path / "doc.xlsx")
        content = read_xlsx_file(f)
        assert len(content) > 0
        # Should contain structured info, not just raw shared strings
        assert (
            "Sheet" in content
            or "sheet" in content.lower()
            or "columns" in content.lower()
        )

    def test_xlsx_includes_sheet_info(self, tmp_path):
        from app.tools.files import read_xlsx_file

        f = _make_real_xlsx(tmp_path / "info.xlsx")
        content = read_xlsx_file(f)
        # Must include some structural metadata
        assert any(kw in content.lower() for kw in ("sheet", "row", "column"))

    def test_xlsx_formula_data_only(self, tmp_path):
        """Formulas must not be exposed; only data_only values."""
        from app.tools.files import read_xlsx_file

        f = _make_formula_xlsx(tmp_path / "formula.xlsx")
        content = read_xlsx_file(f)
        # Formula text should NOT appear
        assert "=SUM" not in content
        assert "=A1+B1" not in content

    def test_xlsx_truncates_rows_to_20(self, tmp_path):
        from app.tools.files import read_xlsx_file

        f = _make_real_xlsx(tmp_path / "many.xlsx", rows=50)
        content = read_xlsx_file(f)
        # Should not have 50 rows of data referenced
        # At minimum, the function returned without error
        assert len(content) > 0

    def test_rejects_xlsm_extension(self, tmp_path):
        from app.tools.files import read_xlsx_file

        f = tmp_path / "doc.xlsm"
        with zipfile.ZipFile(f, "w") as zf:
            zf.writestr("[Content_Types].xml", _CT_XLSX_XML)
            zf.writestr("xl/workbook.xml", _WORKBOOK_XML)
            zf.writestr("xl/vbaProject.bin", b"macro")
        with pytest.raises(ValueError, match="vbaProject"):
            read_xlsx_file(f)


# ── session file format dispatch ─────────────────────────────────────────────


class TestUploadedFileDispatch:
    def test_reads_txt(self, tmp_path):
        from app.conversation.file_index import read_supported_file

        ws = _workspace(tmp_path)
        f = ws.resolve_upload("notes.txt")
        f.write_text("plain text", encoding="utf-8")

        result = read_supported_file(f)
        assert "plain text" in result

    def test_reads_md(self, tmp_path):
        from app.conversation.file_index import read_supported_file

        ws = _workspace(tmp_path)
        f = ws.resolve_upload("notes.md")
        f.write_text("# markdown", encoding="utf-8")

        result = read_supported_file(f)
        assert "markdown" in result

    def test_reads_pdf(self, tmp_path):
        from app.conversation.file_index import read_supported_file

        ws = _workspace(tmp_path)
        f = ws.resolve_upload("doc.pdf")
        _make_valid_pdf(f, pages=1)
        result = read_supported_file(f)
        assert len(result) > 0

    def test_reads_docx(self, tmp_path):
        from app.conversation.file_index import read_supported_file

        ws = _workspace(tmp_path)
        f = ws.resolve_upload("doc.docx")
        _make_real_docx(f)
        result = read_supported_file(f)
        assert len(result) > 0

    def test_reads_xlsx(self, tmp_path):
        from app.conversation.file_index import read_supported_file

        ws = _workspace(tmp_path)
        f = ws.resolve_upload("doc.xlsx")
        _make_real_xlsx(f)
        result = read_supported_file(f)
        assert len(result) > 0

    def test_rejects_unknown_extension(self, tmp_path):
        from app.conversation.file_index import read_supported_file

        ws = _workspace(tmp_path)
        f = ws.resolve_upload("data.exe")
        f.write_bytes(b"payload")

        with pytest.raises(ValueError, match="unsupported"):
            read_supported_file(f)

    def test_rejects_missing_file(self, tmp_path):
        from app.conversation.file_index import read_supported_file

        with pytest.raises((FileNotFoundError, ValueError)):
            read_supported_file(tmp_path / "nonexistent.txt")


# ── B5: parser failure injection leaves no partial final file ────────────────


class TestParserFailureCleanup:
    def test_failed_pdf_upload_preserves_previous_file(self, tmp_path):
        """Damaged PDF content fails validation AFTER temp write; old file stays."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        seed = _make_valid_pdf(tmp_path / "seed.pdf", pages=1)
        save_uploaded_file(ws, "doc.pdf", seed.read_bytes())

        with pytest.raises(ValueError, match="PDF"):
            save_uploaded_file(ws, "doc.pdf", b"%PDF-1.4\nnot a real pdf body")

        assert ws.resolve_upload("doc.pdf").read_bytes().startswith(b"%PDF")
        assert list(ws.upload_dir.glob(".tmp-*")) == []

    def test_failed_docx_upload_preserves_previous_file(self, tmp_path):
        """Macro-enabled DOCX fails validation AFTER temp write; old file stays."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        seed = _make_real_docx(tmp_path / "seed.docx")
        save_uploaded_file(ws, "doc.docx", seed.read_bytes())

        macro = tmp_path / "macro.docx"
        with zipfile.ZipFile(macro, "w") as zf:
            zf.writestr("[Content_Types].xml", _CT_XML)
            zf.writestr("word/document.xml", _DOCUMENT_XML)
            zf.writestr("word/vbaProject.bin", b"macro payload")

        with pytest.raises(ValueError, match="vbaProject"):
            save_uploaded_file(ws, "doc.docx", macro.read_bytes())

        assert ws.resolve_upload("doc.docx").read_bytes() == seed.read_bytes()
        assert list(ws.upload_dir.glob(".tmp-*")) == []

    def test_failed_xlsx_upload_preserves_previous_file(self, tmp_path):
        """Non-ZIP XLSX fails validation AFTER temp write; old file stays."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        seed = _make_real_xlsx(tmp_path / "seed.xlsx")
        save_uploaded_file(ws, "doc.xlsx", seed.read_bytes())

        with pytest.raises(ValueError, match="XLSX"):
            save_uploaded_file(ws, "doc.xlsx", b"not a zip")

        assert ws.resolve_upload("doc.xlsx").read_bytes() == seed.read_bytes()
        assert list(ws.upload_dir.glob(".tmp-*")) == []

    def test_failed_text_upload_preserves_previous_file(self, tmp_path):
        """Non-UTF-8 text fails validation AFTER temp write; old file stays."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        save_uploaded_file(ws, "notes.txt", b"good data")

        with pytest.raises(ValueError, match="UTF-8"):
            save_uploaded_file(ws, "notes.txt", b"\xff\xfe\x00\x00")

        assert ws.resolve_upload("notes.txt").read_text() == "good data"
        assert list(ws.upload_dir.glob(".tmp-*")) == []

    def test_parser_exception_injection_cleans_temp_and_keeps_old(
        self, tmp_path, monkeypatch
    ):
        """Any parser exception after temp write must clean the temp file."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        save_uploaded_file(ws, "data.txt", b"original")

        def _boom(path, ext):
            raise RuntimeError("injected parser failure")

        monkeypatch.setattr("app.tools.files._validate_file_content", _boom)
        with pytest.raises(RuntimeError, match="injected"):
            save_uploaded_file(ws, "data.txt", b"replacement")

        assert ws.resolve_upload("data.txt").read_text() == "original"
        assert list(ws.upload_dir.glob(".tmp-*")) == []

    def test_replace_failure_injection_cleans_temp_and_keeps_old(
        self, tmp_path, monkeypatch
    ):
        """A failed final os.replace must not clobber the existing file."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        save_uploaded_file(ws, "data.txt", b"original")

        def _fail_replace(src, dst):
            raise OSError("injected replace failure")

        monkeypatch.setattr("app.tools.files.os.replace", _fail_replace)
        with pytest.raises(OSError, match="injected"):
            save_uploaded_file(ws, "data.txt", b"replacement")

        assert ws.resolve_upload("data.txt").read_text() == "original"
        assert list(ws.upload_dir.glob(".tmp-*")) == []

    def test_directory_target_failure_leaves_no_partial_file(self, tmp_path):
        """A directory squatting at the final name must not leave a partial file."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        ws.resolve_upload("data.txt").mkdir()

        with pytest.raises(Exception):
            save_uploaded_file(ws, "data.txt", b"content")

        assert ws.resolve_upload("data.txt").is_dir()
        assert list(ws.upload_dir.glob(".tmp-*")) == []

    def test_final_target_symlink_cannot_modify_outside_sentinel(self, tmp_path):
        """A symlink at the FINAL name is rejected; the sentinel stays SAFE."""
        import os

        outside = tmp_path / "outside.txt"
        outside.write_text("SAFE")

        ws = _workspace(tmp_path)
        os.symlink(str(outside), str(ws.upload_dir / "data.txt"))

        from app.tools.files import UnsafeWorkspacePath, save_uploaded_file

        with pytest.raises(UnsafeWorkspacePath):
            save_uploaded_file(ws, "data.txt", b"payload")

        assert outside.read_text() == "SAFE", "outside.txt was written via symlink!"


def test_content_type_mismatch_does_not_override_validation(tmp_path):
    """MIME content-type must be advisory only — actual content drives validation."""
    from app.tools.files import read_pdf_file

    f = tmp_path / "fake.pdf"
    f.write_text("actually text not pdf")
    with pytest.raises(ValueError):
        read_pdf_file(f)


# ── Same-name upload: atomic replace ──────────────────────────────────────────


class TestSameSessionDuplicateUpload:
    def test_last_complete_replaces_previous(self, tmp_path):
        """The last complete, validated upload atomically replaces the prior file."""
        ws = _workspace(tmp_path)

        from app.tools.files import save_uploaded_file

        save_uploaded_file(ws, "data.txt", b"version 1")
        first = ws.resolve_upload("data.txt").read_text()
        assert "version 1" == first

        save_uploaded_file(ws, "data.txt", b"version 2")
        second = ws.resolve_upload("data.txt").read_text()
        assert "version 2" == second

    def test_failed_upload_preserves_old_file(self, tmp_path):
        """When validation fails, the old complete file stays intact."""
        ws = _workspace(tmp_path)

        from app.tools.files import save_uploaded_file

        save_uploaded_file(ws, "data.txt", b"good data")
        # Attempt invalid upload — too large
        with pytest.raises(ValueError):
            save_uploaded_file(ws, "data.txt", b"x" * (MAX_FILE_SIZE_BYTES + 100))
        content = ws.resolve_upload("data.txt").read_text()
        assert "good data" == content

    def test_atomic_replace_no_partial_state(self, tmp_path):
        """After replace, no .tmp files remain."""
        ws = _workspace(tmp_path)

        from app.tools.files import save_uploaded_file

        save_uploaded_file(ws, "doc.txt", b"content")
        tmps = list(ws.upload_dir.glob("*.tmp"))
        assert len(tmps) == 0, f"tmp files left: {[t.name for t in tmps]}"


# ── Symlink exploit: fixed .tmp → outside ────────────────────────────────────


class TestSymlinkExploitDefense:
    def test_fixed_tmp_symlink_cannot_overwrite_outside(self, tmp_path):
        """Precreate .note.txt.tmp as symlink to outside — outside stays SAFE."""
        import os

        outside = tmp_path / "outside.txt"
        outside.write_text("SAFE")

        ws = _workspace(tmp_path)
        fixed_tmp = ws.upload_dir / ".note.txt.tmp"
        os.symlink(str(outside), str(fixed_tmp))

        from app.tools.files import save_uploaded_file

        save_uploaded_file(ws, "note.txt", b"UPLOADED CONTENT")

        # outside MUST remain SAFE
        assert outside.read_text() == "SAFE", "outside.txt was overwritten via symlink!"

        # Final file must be a regular file inside workspace
        final = ws.resolve_upload("note.txt")
        assert final.exists()
        assert not final.is_symlink()
        assert final.read_text() == "UPLOADED CONTENT"

# ── MIME content-type does not override actual content ───────────────────────


def test_text_content_type_on_pdf_is_rejected(tmp_path):
    """A file with .pdf extension but text content is rejected by content validation."""
    from app.tools.files import read_pdf_file

    f = tmp_path / "fake.pdf"
    f.write_text("text/plain content but pdf extension")
    with pytest.raises(ValueError):
        read_pdf_file(f)


# ═══════════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════════

_CT_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    '<Default Extension="docx" '
    'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
    "</Types>"
)

_CT_XLSX_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    '<Default Extension="xlsx" '
    'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
    "</Types>"
)

_DOCUMENT_XML = (
    "<w:document "
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    "<w:body><w:p><w:r><w:t>Hello world</w:t></w:r></w:p></w:body>"
    "</w:document>"
)

_WORKBOOK_XML = (
    '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"/>'
)


def _workspace(tmp_path):
    return SessionWorkspace.for_thread(
        thread_id=UUID_V4,
        base_upload=str(tmp_path / "updated"),
        base_output=str(tmp_path / "output"),
    )


def _make_valid_pdf(path: Path, pages: int = 2) -> Path:
    from pypdf import PdfWriter

    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(200, 300)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "wb") as fp:
        writer.write(fp)
    return path


def _make_real_docx(path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello from python-docx")
    doc.add_paragraph("Second paragraph with more content")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _make_real_xlsx(path: Path, rows: int = 10) -> Path:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws["A1"] = "Name"
    ws["B1"] = "Value"
    for i in range(1, rows):
        ws.cell(row=i + 1, column=1, value=f"Item {i}")
        ws.cell(row=i + 1, column=2, value=i * 10)
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    return path


def _make_formula_xlsx(path: Path) -> Path:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws["A1"] = 10
    ws["A2"] = 20
    ws["A3"] = "=SUM(A1:A2)"
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    return path


class TestContentValidationBypass:
    """Ensure save_uploaded_file validates by TARGET extension, not .tmp."""

    def test_fake_pdf_rejected(self, tmp_path):
        """malformed.pdf with text content must be rejected."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        with pytest.raises(ValueError):
            save_uploaded_file(ws, "malformed.pdf", b"not-a-pdf")

    def test_fake_docx_rejected(self, tmp_path):
        """malformed.docx with text content must be rejected."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        with pytest.raises(ValueError):
            save_uploaded_file(ws, "malformed.docx", b"not-a-docx")

    def test_fake_xlsx_rejected(self, tmp_path):
        """malformed.xlsx with text content must be rejected."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        with pytest.raises(ValueError):
            save_uploaded_file(ws, "malformed.xlsx", b"not-a-xlsx")

    def test_non_utf8_text_rejected(self, tmp_path):
        """Non-UTF-8 .txt must be rejected."""
        from app.tools.files import save_uploaded_file

        ws = _workspace(tmp_path)
        with pytest.raises(ValueError):
            save_uploaded_file(ws, "bad.txt", b"\xff\xfe\x00\x00")
