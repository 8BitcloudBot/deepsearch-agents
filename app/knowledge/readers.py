"""Safe workspace paths and file readers.

Path traversal is REJECTED (not silently sanitized). Uses real
pypdf, python-docx, openpyxl — no regex-only page counting.
Macro and ZIP bomb defense included.
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

# ── Constants ─────────────────────────────────────────────────────────────────

UUID_RE = re.compile(r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$")
ALLOWED_EXTENSIONS = frozenset({".txt", ".md", ".pdf", ".docx", ".xlsx"})
DISALLOWED_EXTENSIONS = frozenset({".doc", ".xls", ".docm", ".xlsm"})
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MiB
MAX_TEXT_CHARS = 100_000
MAX_ZIP_ENTRIES = 500
MAX_UNCOMPRESSED_SIZE = 100 * 1024 * 1024  # 100 MiB
MAX_ZIP_RATIO = 200
MAX_XLSX_DATA_ROWS = 200  # 每个工作表入库的数据行上限（G5：20 行静默丢失实测不可用）
TRUNCATION_SUFFIX = "\n\n[内容已截断：超过最大允许字符数]\n"

MACRO_ENTRIES = frozenset(
    {
        "vbaProject.bin",
        "word/vbaProject.bin",
        "xl/vbaProject.bin",
        "word/vbaData.xml",
        "xl/vbaData.xml",
    }
)
MACRO_CONTENT_TYPES = frozenset(
    {
        "application/vnd.ms-word.document.macroEnabled.main+xml",
        "application/vnd.ms-excel.sheet.macroEnabled.main+xml",
    }
)


# ── Exceptions ────────────────────────────────────────────────────────────────


class UnsafeWorkspacePath(ValueError):  # noqa: N818
    """A path attempted to escape the workspace boundary."""


# ── SessionWorkspace ──────────────────────────────────────────────────────────


class SessionWorkspace:
    """Thread-scoped upload and output directories.

    The only blessed creation path is for_thread() which validates
    the server-assigned UUID. Direct __init__ is allowed but does
    not create directories — callers should prefer for_thread().
    """

    def __init__(self, thread_id: str, base_upload: str, base_output: str):
        self._thread_id = thread_id
        self.upload_dir = Path(base_upload) / f"session_{thread_id}"
        self.output_dir = Path(base_output) / f"session_{thread_id}"

    @classmethod
    def for_thread(cls, *, thread_id: str, base_upload: str, base_output: str):
        if not UUID_RE.fullmatch(thread_id):
            raise ValueError(f"thread_id must be a UUID: {thread_id!r}")
        ws = cls(thread_id, base_upload, base_output)
        ws.upload_dir.mkdir(parents=True, exist_ok=True)
        ws.output_dir.mkdir(parents=True, exist_ok=True)
        return ws

    def _safe_resolve(self, base: Path, name: str) -> Path:
        """Resolve a name within base. REJECT traversal, don't sanitize."""
        if not name or not name.strip():
            raise UnsafeWorkspacePath("Empty filename not allowed")

        # Reject absolute paths
        if os.path.isabs(name) or name.startswith("/"):
            raise UnsafeWorkspacePath(f"Absolute path not allowed: {name!r}")

        # Reject Windows backslash traversal
        if "\\" in name:
            raise UnsafeWorkspacePath(f"Backslash not allowed: {name!r}")

        # Reject directory components — only basename allowed
        clean = Path(name)
        if str(clean) != clean.name or clean.name in (".", ".."):
            raise UnsafeWorkspacePath(f"Directory traversal not allowed: {name!r}")
        if not clean.name or clean.name in (".", ".."):
            raise UnsafeWorkspacePath(f"Invalid filename: {name!r}")

        candidate = (base / clean.name).resolve()
        base_resolved = base.resolve()

        # Use is_relative_to for proper containment
        if not candidate.is_relative_to(base_resolved):
            raise UnsafeWorkspacePath(f"Path escape attempt: {name!r}")

        # Symlink check
        try:
            real = candidate.resolve(strict=False)
            if not real.is_relative_to(base_resolved):
                raise UnsafeWorkspacePath(f"Symlink escape: {name!r}")
        except (OSError, RuntimeError):
            raise UnsafeWorkspacePath(f"Cannot resolve path: {name!r}")

        return candidate

    def resolve_upload(self, name: str) -> Path:
        return self._safe_resolve(self.upload_dir, name)

    def resolve_output(self, name: str) -> Path:
        return self._safe_resolve(self.output_dir, name)


# ── Atomic write helper ───────────────────────────────────────────────────────


def _write_all(fd: int, data: bytes) -> None:
    """Write all bytes to fd, handling short writes from os.write."""
    written = 0
    while written < len(data):
        n = os.write(fd, data[written:])
        if n == 0:
            raise OSError("os.write returned 0")
        written += n


def _atomic_write_bytes(target: Path, data: bytes) -> None:
    """Write data via random exclusive temp file, fsync, then os.replace.

    Uses mkstemp (O_EXCL) to reject pre-existing temp symlinks.
    Handles short writes via _write_all.
    """
    target.parent.mkdir(parents=True, exist_ok=True)
    fd = -1
    tmp_path = None
    try:
        fd, tmp_path = tempfile.mkstemp(
            dir=str(target.parent), prefix=".tmp-", suffix=".tmp"
        )
        _write_all(fd, data)
        os.fsync(fd)
        os.close(fd)
        fd = -1
        os.replace(tmp_path, target)
    finally:
        if fd >= 0:
            os.close(fd)
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


def _atomic_write_and_validate(target: Path, data: bytes, validator, ext: str) -> None:
    """Write via random exclusive temp file, validate, then os.replace.

    validator(path, ext) receives the FINAL target extension (not .tmp).
    This prevents content-type bypass where e.g. a .pdf file's content
    would be validated as .tmp and bypass PDF header checks.
    """
    target.parent.mkdir(parents=True, exist_ok=True)
    fd = -1
    tmp_path = None
    try:
        fd, tmp_path = tempfile.mkstemp(
            dir=str(target.parent), prefix=".tmp-", suffix=".tmp"
        )
        _write_all(fd, data)
        os.fsync(fd)
        os.close(fd)
        fd = -1
        validator(Path(tmp_path), ext)
        os.replace(tmp_path, target)
        tmp_path = None  # consumed by replace
    finally:
        if fd >= 0:
            os.close(fd)
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ── Upload helpers ────────────────────────────────────────────────────────────


def save_uploaded_file(workspace: SessionWorkspace, name: str, data: bytes) -> Path:
    """Atomically save an uploaded file with validation.

    1. Reject names with directory components.
    2. Validate extension.
    3. Check size before writing.
    4. Write to temp file, validate content, then atomic replace.
    5. On validation failure, delete temp, old file remains.
    """
    # Reject directory components upfront
    if "/" in name or "\\" in name:
        raise UnsafeWorkspacePath(f"Directory component not allowed: {name!r}")

    resolved = workspace.resolve_upload(name)

    # Check extension
    ext = resolved.suffix.lower()
    if ext in DISALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型：{ext}")
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件类型：{ext}（支持：{', '.join(sorted(ALLOWED_EXTENSIONS))}）"
        )

    # Check size before writing
    if len(data) > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"文件过大：{len(data)} 字节（上限 {MAX_FILE_SIZE_BYTES}）"
        )

    # Write temp, validate with final extension, atomic replace
    _atomic_write_and_validate(resolved, data, _validate_file_content, ext)
    return resolved


def validate_upload_file(path: Path) -> None:
    """Validate a single uploaded file for size and extension."""
    ext = path.suffix.lower()
    if ext in DISALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型：{ext}")
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件类型：{ext}（支持：{', '.join(sorted(ALLOWED_EXTENSIONS))}）"
        )
    if path.stat().st_size > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"文件过大：{path.stat().st_size} 字节（上限 {MAX_FILE_SIZE_BYTES}）"
        )


def _validate_file_content(path: Path, ext: str) -> None:
    """Validate content based on the FINAL target extension.

    ext is the extension of the target file (e.g. ".pdf"), NOT the
    temp file's suffix (which is always ".tmp"). This prevents
    content-type bypass via temp suffix confusion.
    """
    if ext == ".pdf":
        read_pdf_file(path)
    elif ext == ".docx":
        read_docx_file(path)
    elif ext == ".xlsx":
        read_xlsx_file(path)
    elif ext in (".txt", ".md"):
        read_text_file(path)


# ── Text readers ──────────────────────────────────────────────────────────────


def read_text_file(path: Path, max_chars: int = MAX_TEXT_CHARS) -> str:
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise ValueError("文件不是有效的 UTF-8 编码，请另存为 UTF-8 后重试") from exc
    if len(raw) > max_chars:
        return raw[:max_chars] + TRUNCATION_SUFFIX
    return raw


# ── PDF reader (pypdf) ────────────────────────────────────────────────────────


def read_pdf_file(path: Path, max_pages: int = 200) -> str:
    """Read PDF using pypdf with real page counting and text extraction."""
    data = path.read_bytes()
    if not data.startswith(b"%PDF"):
        raise ValueError("不是有效的 PDF 文件（缺少 %PDF 头）")

    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf not installed. Run: uv sync --extra dev") from exc

    try:
        reader = PdfReader(path)
    except Exception:
        raise ValueError("PDF 文件损坏或无法读取")

    if reader.is_encrypted:
        raise ValueError("暂不支持加密 PDF")

    pages = len(reader.pages)
    if pages > max_pages:
        raise ValueError(f"PDF 页数超出上限：{pages} > {max_pages}")

    texts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text()
            if t:
                texts.append(t)
        except Exception:
            pass

    combined = "\n".join(texts)
    if len(combined) > MAX_TEXT_CHARS:
        combined = combined[:MAX_TEXT_CHARS] + TRUNCATION_SUFFIX
    if not combined:
        combined = "[PDF 未包含可提取的文本]"
    return combined


# ── DOCX reader (python-docx) ─────────────────────────────────────────────────


def _validate_zip_safety(zf: zipfile.ZipFile) -> None:
    """Reject ZIP bombs and macro-enabled archives."""
    names = zf.namelist()
    if len(names) > MAX_ZIP_ENTRIES:
        raise ValueError(f"ZIP 条目数超出上限：{len(names)} > {MAX_ZIP_ENTRIES}")

    # Check for macro entries
    for name in names:
        base = name.split("/")[-1]
        if base in MACRO_ENTRIES or name in MACRO_ENTRIES:
            raise ValueError("已拒绝宏文档：检测到 vbaProject.bin")

    # Check uncompressed size to avoid ZIP bomb
    total_size = 0
    for info in zf.infolist():
        if info.file_size > 0:
            total_size += info.file_size
            # Check per-entry ratio
            if info.compress_size > 0:
                ratio = info.file_size / info.compress_size
                if ratio > MAX_ZIP_RATIO:
                    raise ValueError("检测到疑似 ZIP 炸弹：异常压缩比")
    if total_size > MAX_UNCOMPRESSED_SIZE:
        raise ValueError(
            f"ZIP 解压后大小超出上限：{total_size} > {MAX_UNCOMPRESSED_SIZE}"
        )


def _check_content_types_macros(names: list[str], raw_ct: bytes | None) -> None:
    """Reject macro content types in [Content_Types].xml."""
    if raw_ct:
        text = raw_ct.decode("utf-8", errors="replace").lower()
        for mt in MACRO_CONTENT_TYPES:
            if mt.lower() in text:
                raise ValueError("检测到宏启用内容类型")


def read_docx_file(path: Path) -> str:
    """Read DOCX using python-docx with macro and ZIP bomb defense."""
    try:
        import zipfile as _zf

        with _zf.ZipFile(path) as zf:
            names = zf.namelist()
            if "[Content_Types].xml" not in names:
                raise ValueError("DOCX 缺少 [Content_Types].xml")
            if "word/document.xml" not in names:
                raise ValueError("DOCX 缺少 word/document.xml")

            # Check content types for macros
            ct_data = zf.read("[Content_Types].xml")
            _check_content_types_macros(names, ct_data)

            # ZIP bomb defense
            _validate_zip_safety(zf)
    except zipfile.BadZipFile as exc:
        raise ValueError("不是有效的 DOCX 文件（ZIP 结构损坏）") from exc
    except ValueError:
        raise

    try:
        from docx import Document

        doc = Document(str(path))
    except ImportError as exc:
        raise RuntimeError(
            "python-docx not installed. Run: uv sync --extra dev"
        ) from exc
    except Exception:
        raise ValueError("无法解析 DOCX 文档")

    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))

    text = "\n".join(paragraphs)
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + TRUNCATION_SUFFIX
    return text


# ── XLSX reader (openpyxl) ────────────────────────────────────────────────────


def read_xlsx_file(path: Path) -> str:
    """Read XLSX using openpyxl with read_only, data_only, no macros."""
    try:
        import zipfile as _zf

        with _zf.ZipFile(path) as zf:
            names = zf.namelist()
            if "[Content_Types].xml" not in names:
                raise ValueError("XLSX 缺少 [Content_Types].xml")
            if "xl/workbook.xml" not in names:
                raise ValueError("XLSX 缺少 xl/workbook.xml")

            ct_data = zf.read("[Content_Types].xml")
            _check_content_types_macros(names, ct_data)
            _validate_zip_safety(zf)
    except zipfile.BadZipFile as exc:
        raise ValueError("不是有效的 XLSX 文件（ZIP 结构损坏）") from exc
    except ValueError:
        raise

    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl not installed. Run: uv sync --extra dev") from exc

    import itertools

    output_lines: list[str] = []
    try:
        # openpyxl rejects paths whose extension is not a supported format
        # (e.g. the ".tmp" of a validation temp file), so hand it a binary
        # file object: content parsing stays extension-independent, matching
        # the documented "parse content, not names" contract.
        with path.open("rb") as fh:
            wb = load_workbook(fh, read_only=True, data_only=True, keep_links=False)
            try:
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    # Bound: header + up to MAX_XLSX_DATA_ROWS data rows
                    all_rows = ws.iter_rows(values_only=True)
                    header_row = next(all_rows, None)
                    headers = (
                        [str(c) if c is not None else "" for c in header_row]
                        if header_row
                        else []
                    )
                    data_rows = list(itertools.islice(all_rows, MAX_XLSX_DATA_ROWS))
                    truncated = next(all_rows, None) is not None

                    col_count = ws.max_column or len(headers)
                    row_count = len(data_rows) + (1 if header_row else 0)

                    output_lines.append(f"## Sheet: {sheet_name}")
                    output_lines.append(
                        f"Rows (bounded): {row_count}, Columns: {col_count}"
                    )
                    if headers:
                        output_lines.append(f"Headers: {' | '.join(headers)}")
                    output_lines.append("")
                    for row in data_rows:
                        row_str = " | ".join(
                            str(c) if c is not None else "" for c in row
                        )
                        output_lines.append(row_str)
                    if truncated:
                        output_lines.append(
                            f"[数据已截断：仅入库每个工作表前 {MAX_XLSX_DATA_ROWS} 行]"
                        )
                    output_lines.append("")
            finally:
                wb.close()
    except Exception:
        raise ValueError("无法解析 XLSX 工作簿")
    text = "\n".join(output_lines)
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + TRUNCATION_SUFFIX
    return text
